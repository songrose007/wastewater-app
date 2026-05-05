"""Generate DOCX report for Taizhou medical wastewater project."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    return h


def add_p(text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p


def add_tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ''
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def page_break():
    doc.add_page_break()


# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph()

add_p('泰州20t/d医疗机构废水处理工程', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p('工艺设计方案', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

info = [
    '工程名称：泰州医疗机构废水处理工程',
    '设计规模：20 m3/d',
    '废水类型：医疗机构废水（医院污水）',
    '排放标准：GB18466-2005 预处理标准',
    '主体工艺：A/O生化+沉淀+消毒',
    '工程投资：约 16.45 万元',
    '编制日期：%s' % datetime.date.today().strftime('%Y年%m月'),
]
for line in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

page_break()

# ===== Ch1 =====
add_h('第1章 工程概况', 1)
add_h('1.1 项目背景', 2)
add_p('本项目为泰州地区某医疗机构配套废水处理工程，设计处理能力为20 m3/d。出水执行GB18466-2005预处理标准，达标后排入市政污水管网。')

add_h('1.2 工程规模', 2)
add_tbl(
    ['项目', '参数', '备注'],
    [['设计水量', '20 m3/d', '日均处理能力'],
     ['时均水量', '0.83 m3/h', '按24h运行'],
     ['峰值系数', '1.5', 'Kz=1.5'],
     ['峰值水量', '1.25 m3/h', ''],
     ['运行模式', '连续运行', ''],
     ['占地面积', '约40 m2', '不含利旧设备房']])

# ===== Ch2 =====
page_break()
add_h('第2章 设计依据与标准', 1)
add_h('2.1 设计规范', 2)
for n in [
    'GB18466-2005 医疗机构水污染物排放标准',
    'HJ2029-2013 医院污水处理工程技术规范',
    'GB50014-2021 室外排水设计规范',
    'GB50015-2019 建筑给水排水设计规范',
    'GB50069-2002 给水排水工程构筑物结构设计规范',
    'GB8978-1996 污水综合排放标准',
]:
    add_p(n)

add_h('2.2 出水水质要求', 2)
add_tbl(
    ['控制项目', '预处理标准限值', '单位', '参考依据'],
    [['CODcr', '250', 'mg/L', 'GB18466-2005 表2'],
     ['BOD5', '100', 'mg/L', 'GB18466-2005 表2'],
     ['SS', '60', 'mg/L', 'GB18466-2005 表2'],
     ['NH3-N', '不考核', 'mg/L', '排入城市管网'],
     ['粪大肠菌群数', '5000', 'MPN/L', 'GB18466-2005 表2'],
     ['pH', '6~9', '', 'GB18466-2005 表2']])

# ===== Ch3 =====
page_break()
add_h('第3章 设计水量与水质', 1)
add_h('3.1 设计进水水质', 2)
add_p('参考同类医疗机构废水水质及规范推荐值：')
add_tbl(
    ['控制项目', '进水浓度', '单位', '来源依据'],
    [['CODcr', '350', 'mg/L', 'HJ2029-2013 典型值'],
     ['BOD5', '150', 'mg/L', 'HJ2029-2013 典型值'],
     ['SS', '200', 'mg/L', 'HJ2029-2013 典型值'],
     ['NH3-N', '40', 'mg/L', '规范推荐值'],
     ['粪大肠菌群数', '1.6x10^8', 'MPN/L', '规范推荐值'],
     ['pH', '6~9', '', '']])

add_h('3.2 污染物去除率分析', 2)
add_tbl(
    ['控制项目', '进水', '出水要求', '去除率', '可满足性'],
    [['CODcr', '350 mg/L', '250 mg/L', '28.6%', 'A/O工艺去除率>70%, 满足'],
     ['BOD5', '150 mg/L', '100 mg/L', '33.3%', 'A/O工艺去除率>85%, 满足'],
     ['SS', '200 mg/L', '60 mg/L', '70%', '沉淀去除率>80%, 满足']])

# ===== Ch4 =====
page_break()
add_h('第4章 处理工艺选择', 1)
add_h('4.1 工艺选择原则', 2)
add_p('医疗机构废水处理工艺选择应根据废水性质、排放标准及场地条件等因素综合确定。本项目为常规医疗废水（非传染病医院），要求达到预处理标准后排入市政管网，故选择以生化处理为核心的工艺路线。')

add_h('4.2 推荐工艺流程', 2)
add_p('医院废水 -> 格栅井 -> 一体化A/O生化池（缺氧区+好氧区+沉淀区）-> 消毒接触池（次氯酸钠消毒）-> 达标排放', bold=True)

add_h('4.3 工艺特点', 2)
for f in [
    '一体化玻璃钢罐体集成缺氧-好氧-沉淀三区，结构紧凑，占地小',
    '内循环回流（混合液+污泥）确保脱氮除碳效果',
    '次氯酸钠消毒，安全可靠，无液氯泄漏风险',
    'PLC自控运行，管理简单',
    '玻璃钢材质耐腐蚀，使用寿命20年以上',
    '设备房利旧，土建仅需基础，工程周期短',
]:
    add_p(f)

# ===== Ch5 =====
page_break()
add_h('第5章 主要构筑物及设备', 1)
add_h('5.1 一体化玻璃钢水池', 2)
add_tbl(
    ['参数', '数值'],
    [['规格尺寸', '8.0m x 2.5m x 2.5m'],
     ['材质', '缠丝玻璃钢（FRP），含人孔'],
     ['有效容积', '约40 m3'],
     ['HRT', '约48h'],
     ['功能分区', '缺氧区+好氧区+沉淀区']])

add_h('5.2 主要设备清单', 2)
add_tbl(
    ['序号', '设备名称', '规格型号', '数量', '单位', '单价(元)', '金额(元)'],
    [['1', '一体化玻璃钢水池', '8.0x2.5x2.5m,FRP', '1', '座', '88000', '88000'],
     ['2', '混合曝气系统', 'DN50穿孔曝气管', '1', '套', '3000', '3000'],
     ['3', '提升泵', 'Q=2m3/h,H=10m,0.55kW', '2', '台', '1200', '2400'],
     ['4', '浮球液位计', '0~3m', '1', '只', '300', '300'],
     ['5', '微孔曝气器', 'BNQZ-192,15只', '1', '套', '3000', '3000'],
     ['6', '中心筒', '400x800mm', '1', '台', '1000', '1000'],
     ['7', '污泥回流泵', 'Q=2m3/h,H=10m,0.55kW', '2', '台', '1200', '2400'],
     ['8', '混合液回流泵', 'Q=5m3/h,H=10m,0.75kW', '2', '台', '1400', '2800'],
     ['9', '消毒加药装置', '1m3药桶+搅拌+计量泵x2', '1', '套', '4500', '4500'],
     ['10', '回转式风机', 'HC30S', '2', '台', '4500', '9000'],
     ['11', '电气自控柜', 'PLC系统', '1', '套', '22000', '22000'],
     ['12', '管阀件', '', '1', '套', '1500', '1500']])

# ===== Ch6 =====
page_break()
add_h('第6章 工艺设计参数', 1)
add_h('6.1 主要设计参数', 2)
add_tbl(
    ['构筑物/工艺段', '设计参数', '数值', '单位'],
    [['一体化A/O池', '总有效容积', '40', 'm3'],
     ['', 'HRT', '48', 'h'],
     ['', '缺氧区容积', '12', 'm3'],
     ['', '好氧区容积', '28', 'm3'],
     ['好氧区', 'MLSS', '3000-4000', 'mg/L'],
     ['', 'BOD污泥负荷', '0.08-0.12', 'kgBOD/kgMLSS.d'],
     ['', '气水比', '15:1', ''],
     ['沉淀区', '表面负荷', '0.6-0.8', 'm3/m2.h'],
     ['', '污泥回流比', '50-100%', ''],
     ['消毒', '消毒剂', 'NaClO', ''],
     ['', '有效氯投加量', '20-30', 'mg/L'],
     ['', '接触时间', '1.0', 'h']])

add_h('6.2 曝气量计算', 2)
add_p('BOD5进水: 150 mg/L, BOD5去除量: 3.0 kgBOD/d')
add_p('需氧量: 3.0 x 1.5 = 4.5 kgO2/d')
add_p('氧转移效率(微孔曝气): 15%')
add_p('实际设计气量: 300 m3/d (12.5 m3/h), 气水比: 15:1')

add_h('6.3 污泥产量', 2)
add_p('干污泥产量: 3.0 x 0.4 = 1.2 kg/d')
add_p('剩余污泥湿量(含水率99%): 120 L/d, 浓缩后(含水率97%): 40 L/d')
add_p('污泥处置: 定期抽排，加次氯酸钠消毒后委托有资质单位外运处置')

# ===== Ch7 =====
page_break()
add_h('第7章 运行管理', 1)
add_h('7.1 日常运行控制', 2)
add_tbl(
    ['控制项目', '控制范围', '检测频率', '备注'],
    [['好氧区DO', '2-4 mg/L', '每日1次', '调节风机运行'],
     ['SV30', '15-30%', '每日1次', '判断污泥沉降'],
     ['MLSS', '3000-4000 mg/L', '每周1次', ''],
     ['pH', '6.5-8.5', '每日1次', ''],
     ['NaClO投加量', '20-30 mg/L', '随时', '根据余氯调整'],
     ['出水COD', '250 mg/L', '每周1-2次', '委托检测'],
     ['粪大肠菌群', '5000 MPN/L', '每月1次', '委托检测']])

add_h('7.2 污泥管理', 2)
add_p('医疗机构污泥属于危险废物（HW01）。池底沉积污泥每3-6个月通过排泥泵抽出，投加次氯酸钠或石灰消毒（有效氯2-5g/L污泥），脱水后委托有资质的医疗废物处置单位外运处置。')

# ===== Ch8 =====
page_break()
add_h('第8章 投资估算与经济分析', 1)
add_h('8.1 工程投资', 2)
add_tbl(
    ['序号', '费用类别', '金额(元)', '占比', '备注'],
    [['1', '一体化玻璃钢水池', '88000', '53.5%', '主体设备'],
     ['2', '配套设备', '57900', '35.2%', '曝气/泵/风机/仪表等'],
     ['3', '运输费', '3000', '1.8%', ''],
     ['4', '安装费', '8000', '4.9%', '含吊装、差旅'],
     ['5', '税金(9%增值税)', '13581', '', ''],
     ['', '合计', '164481', '100%', '']])

add_p('工程总投资（含税）: 约 16.45 万元', bold=True)
add_p('吨水投资: 164481 / 20 = 8224 元/m3.d')

add_h('8.2 运行成本分析', 2)
add_tbl(
    ['成本项目', '计算依据', '日费用(元)', '年费用(元)'],
    [['电费', '装机2.5kW,运行系数0.7,电价0.8元/kWh', '33.60', '12264'],
     ['药剂费', 'NaClO 25g/m3 x 20m3/d x 2元/kg', '10.00', '3650'],
     ['污泥处置费', '年产污泥15m3, 500元/m3', '20.55', '7500'],
     ['人工费', '兼职管理, 0.2人x60000元/年', '32.88', '12000'],
     ['合计', '', '97.03', '35414']])

add_p('年运行总成本: 约 3.54 万元/年', bold=True)
add_p('吨水运行成本: 35414 / (20x365) = 4.85 元/m3', bold=True)
add_p('设备折旧(按15年): 145900 / 15 = 9727 元/年')
add_p('吨水综合成本(含折旧): (35414+9727) / (20x365) = 6.18 元/m3', bold=True)

# ===== Appendix =====
page_break()
add_h('附件一  设备清单及报价汇总', 1)
add_tbl(
    ['序号', '设备名称', '规格型号', '数量', '单位', '单价(元)', '金额(元)'],
    [['一', '设备房(利旧)', '', '1', '座', '0', '0'],
     ['二', '设备基础', '9x3x0.3m钢砼', '1', '座', '0', '0'],
     ['1', '一体化玻璃钢水池', '8.0x2.5x2.5m,FRP', '1', '座', '88000', '88000'],
     ['2', '混合曝气系统', 'DN50穿孔曝气管', '1', '套', '3000', '3000'],
     ['3', '提升泵', 'Q=2m3/h,H=10m,0.55kW', '2', '台', '1200', '2400'],
     ['4', '浮球液位计', '0~3m', '1', '只', '300', '300'],
     ['5', '微孔曝气器', 'BNQZ-192,15只', '1', '套', '3000', '3000'],
     ['6', '中心筒', '400x800', '1', '台', '1000', '1000'],
     ['7', '污泥回流泵', 'Q=2m3/h,H=10m,0.55kW', '2', '台', '1200', '2400'],
     ['8', '混合液回流泵', 'Q=5m3/h,H=10m,0.75kW', '2', '台', '1400', '2800'],
     ['9', '消毒加药装置', '1m3药桶+搅拌+计量泵x2', '1', '套', '4500', '4500'],
     ['10', '回转式风机', 'HC30S', '2', '台', '4500', '9000'],
     ['11', '电气自控柜', 'PLC系统', '1', '套', '22000', '22000'],
     ['12', '管阀件', '', '1', '套', '1500', '1500'],
     ['13', '运输费', '', '1', '项', '3000', '3000'],
     ['14', '安装费', '10%(吊装+差旅)', '1', '项', '8000', '8000'],
     ['', '税金(9%增值税)', '', '', '', '', '13581'],
     ['', '总计(含税含安装含调试)', '', '', '', '', '164481']])

page_break()
add_h('附件二  平面布置说明', 1)
add_tbl(
    ['区域', '说明'],
    [['设备房', '利旧，内部布置一体化水池、风机、加药装置、电控柜'],
     ['一体化水池', '8.0x2.5x2.5m玻璃钢罐体，埋地或半地上安装'],
     ['设备基础', '9x3x0.3m钢砼基础，承载力100kPa'],
     ['风机及加药区', '设于设备房内水池一侧，便于管路连接'],
     ['电控柜', '靠墙安装，接入220V/380V电源']])

add_h('布置要点', 2)
for i, pt in enumerate([
    '一体化水池进出水管口方位根据现场条件确定',
    '消毒加药装置靠近水池出水端，减少管道长度',
    '风机设于通风良好处，必要时设隔音罩',
    '电控柜与水池保持安全距离，防止溅水',
    '各泵出口设止回阀和检修阀门，管路最低点设排空阀',
], 1):
    add_p('%d. %s' % (i, pt))

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('--- 方案完 ---')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output_path = r'E:\claude\wastewater-app\backend\reports\taizhou_medical_report.docx'
doc.save(output_path)
print('DOCX saved to:', output_path)
