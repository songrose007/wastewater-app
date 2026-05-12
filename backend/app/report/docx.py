"""DOCX 设计方案生成服务。"""
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any


def generate_docx_from_snapshot(snapshot: dict[str, Any], output_path: str) -> str:
    """根据项目快照生成 DOCX 文件。"""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    style.paragraph_format.line_spacing = 1.5

    def add_run(paragraph, text: str, *, bold: bool = False, size=Pt(11), font="SimSun", color=None):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color:
            run.font.color.rgb = color
        return run

    def heading_1(text: str) -> None:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(18)
        add_run(paragraph, text, bold=True, size=Pt(15), font="SimHei")

    def heading_2(text: str) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        add_run(paragraph, text, bold=True, size=Pt(13), font="SimHei")

    def paragraph(text: str, *, indent: bool = True) -> None:
        item = doc.add_paragraph()
        item.paragraph_format.line_spacing = 1.5
        if indent:
            item.paragraph_format.first_line_indent = Cm(0.74)
        add_run(item, text)

    def table(caption: str, headers: list[str], rows: list[list[Any]]) -> None:
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(cap, caption, bold=True, size=Pt(10), font="SimHei", color=RGBColor(0x1F, 0x4E, 0x79))
        grid = doc.add_table(rows=1 + len(rows), cols=len(headers))
        grid.style = "Table Grid"
        grid.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, header in enumerate(headers):
            cell = grid.rows[0].cells[index]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(header), bold=True, size=Pt(9))
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9D9D9")
            cell._tc.get_or_add_tcPr().append(shading)
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                cell = grid.rows[row_index + 1].cells[col_index]
                cell.text = ""
                add_run(cell.paragraphs[0], "" if value is None else str(value), size=Pt(9))
        doc.add_paragraph()

    project_name = snapshot["project_name"]
    flow_rate = snapshot["flow_rate"]
    route = snapshot["process_route"]
    water_quality = snapshot["water_quality"]
    calculations = snapshot["calculation_results"]
    equipment = snapshot["equipment_list"] or []
    cost = snapshot.get("cost_estimate")

    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, project_name, bold=True, size=Pt(22), font="SimHei")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "污水处理工程设计方案", bold=True, size=Pt(16), font="SimHei")
    for _ in range(3):
        doc.add_paragraph()
    for line in [
        f"设计规模：{flow_rate} m³/d",
        f"废水类型：{snapshot['wastewater_type']}",
        f"排放标准：{snapshot['target_standard']}",
        f"主体工艺：{route.get('route_name_zh', '')}",
        f"编制日期：{date.today().strftime('%Y年%m月')}",
    ]:
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(info, line)
    doc.add_page_break()

    heading_1("第一章  项目概况")
    paragraph(f"本项目为{project_name}，设计处理规模为 {flow_rate} m³/d，出水执行 {snapshot['target_standard']}。")
    table("表1-1  设计进水水质", ["控制项目", "进水浓度", "单位"], [[k, v, _unit_for(k)] for k, v in water_quality.items()])

    heading_1("第二章  工艺方案选择")
    paragraph(f"经工艺比选，推荐采用“{route.get('route_name_zh', route.get('route_id', ''))}”作为本项目主体处理工艺。")
    paragraph(f"该路线综合评分为 {route.get('total_score', 0)}，具体构筑物参数见后续章节。")

    heading_1("第三章  构筑物设计计算")
    for index, result in enumerate(calculations, 1):
        heading_2(f"3.{index}  {result.get('unit_name_zh', result.get('unit_code', ''))}")
        params = result.get("computed_parameters", {})
        rows = [[key, _format_value(value)] for key, value in params.items() if isinstance(value, (int, float, str))]
        if rows:
            table(f"表3-{index}  设计参数", ["参数", "数值"], rows[:30])
        warnings = result.get("warnings") or []
        for warning in warnings:
            paragraph(f"校核提示：{warning}")

    heading_1("第四章  主要设备选型")
    if equipment:
        table(
            "表4-1  主要设备清单",
            ["序号", "类别", "设备名称", "型号", "数量", "单价(元)", "总价(元)", "制造商"],
            [
                [
                    index,
                    item.get("category", ""),
                    item.get("model_name_zh", ""),
                    item.get("model_id", ""),
                    item.get("quantity", 1),
                    f"{item.get('unit_price_cny', 0):,.0f}",
                    f"{item.get('total_price_cny', 0):,.0f}",
                    item.get("manufacturer", ""),
                ]
                for index, item in enumerate(equipment, 1)
            ],
        )
    else:
        paragraph("尚未生成设备选型结果。")

    heading_1("第五章  投资估算与运行成本")
    if cost:
        capex = cost.get("capex", {})
        opex = cost.get("opex", {})
        table(
            "表5-1  工程投资估算",
            ["费用项目", "金额（万元）"],
            [
                ["土建工程费", f"{capex.get('civil_cost', 0):,.2f}"],
                ["设备购置费", f"{capex.get('equipment_cost', 0):,.2f}"],
                ["安装工程费", f"{capex.get('installation_cost', 0):,.2f}"],
                ["设计费", f"{capex.get('engineering_cost', 0):,.2f}"],
                ["不可预见费", f"{capex.get('contingency_cost', 0):,.2f}"],
                ["合计", f"{capex.get('total_capex', 0):,.2f}"],
            ],
        )
        table(
            "表5-2  年运行成本估算",
            ["费用项目", "金额（万元/年）"],
            [
                ["电费", f"{opex.get('energy_cost', 0):,.2f}"],
                ["药剂费", f"{opex.get('chemical_cost', 0):,.2f}"],
                ["人工费", f"{opex.get('labor_cost', 0):,.2f}"],
                ["维护费", f"{opex.get('maintenance_cost', 0):,.2f}"],
                ["污泥处置费", f"{opex.get('sludge_disposal_cost', 0):,.2f}"],
                ["折旧", f"{opex.get('depreciation_cost', 0):,.2f}"],
                ["合计", f"{opex.get('total_annual_opex', 0):,.2f}"],
            ],
        )
        paragraph(f"吨水处理成本约为 {cost.get('cost_per_m3', 0):.2f} 元/m³。", indent=False)
    else:
        paragraph("尚未生成造价估算结果。")

    heading_1("第六章  待确认与复核事项")
    for item in [
        "进水水质应以业主最新检测报告为准。",
        "设备价格为知识库估算价，投标或采购前应结合厂家正式报价复核。",
        "图纸尺寸、标高及既有构筑物条件应由工程师现场复核。",
        "本自动生成方案需由具备相应资质的专业工程师审核后用于正式设计。",
    ]:
        paragraph(item)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "—— 方案完 ——", size=Pt(10), color=RGBColor(0x99, 0x99, 0x99))

    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    return str(target)


def _unit_for(parameter_code: str) -> str:
    return "" if parameter_code.lower() in {"ph", "temperature"} else "mg/L"


def _format_value(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)
