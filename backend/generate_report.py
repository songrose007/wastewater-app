#!/usr/bin/env python3
"""
污水处理设计方案一键生成器
用法: python generate_report.py [project_input.yaml]

输入: YAML 配置文件（水质水量+排放标准+图纸路径+设备清单路径）
输出: DOCX 设计方案（按尚科环境标准模板格式）

依赖: pip install python-docx pillow pyyaml openpyxl pymupdf
"""
import sys, os, yaml, datetime
from pathlib import Path

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent))

from app.knowledge.loader import KnowledgeLoader
from app.engine.process_selector import ProcessSelector
from app.engine.orchestration import CalculationOrchestrator
from app.engine.calculators.registry import CalculatorRegistry
from app.engine.equipment_selector import EquipmentSelector
from app.engine.equipment_verifier import EquipmentVerifier
from app.engine.cost_estimator import CostEstimator
from app.engine.drawing_parser import DrawingParser


def load_input(yaml_path: str) -> dict:
    with open(yaml_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)


def parse_drawings(drawing_paths: dict) -> list:
    """Parse PDF/DWG drawings and extract text elements."""
    parser = DrawingParser()
    elements = []
    for dtype, path in drawing_paths.items():
        if path and os.path.exists(path):
            try:
                result = parser.parse(path)
                for e in result.get('elements', []):
                    e['source'] = dtype
                elements.extend(result.get('elements', []))
                print(f'  [drawing] {dtype}: {len(result.get("elements",[]))} elements from {result.get("page_count",0)} pages')
            except Exception as ex:
                print(f'  [drawing] {dtype}: parse failed - {ex}')
    return elements


def parse_equipment_list(xlsx_path: str) -> list:
    """Parse equipment list from Excel."""
    if not xlsx_path or not os.path.exists(xlsx_path):
        return []
    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        ws = wb[wb.sheetnames[0]]
        rows = []
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, values_only=True):
            rows.append([str(v) if v is not None else '' for v in row])
        print(f'  [equipment] {len(rows)} rows from {wb.sheetnames[0]}')
        return rows
    except Exception as ex:
        print(f'  [equipment] parse failed - {ex}')
        return []


def run_pipeline(config: dict) -> dict:
    """Execute the full design pipeline and return all results."""
    proj = config['project']
    wq = config.get('water_quality', {})
    drawings_cfg = config.get('drawings', {})
    equip_path = config.get('equipment_list', '')
    opts = config.get('options', {})

    flow_m3_d = proj['flow_rate_m3_h'] * 24
    print(f'\n{"="*60}')
    print(f'  项目: {proj["name"]}')
    print(f'  规模: {proj["flow_rate_m3_h"]} m3/h ({flow_m3_d} m3/d)')
    print(f'  类型: {proj["wastewater_type"]}  标准: {proj["target_standard"]}')
    print(f'{"="*60}\n')

    # ---- Init ----
    kb = KnowledgeLoader()
    registry = CalculatorRegistry()

    # ---- Step 1: Process Selection ----
    print('[1/5] 工艺选择...')
    selector = ProcessSelector(kb)
    routes = selector.select(
        wastewater_type=proj['wastewater_type'],
        flow_rate=flow_m3_d,
        design_temp_min=proj.get('design_temp', 20),
        water_quality=wq,
        target_standard_id=proj['target_standard'],
    )
    if not routes.get('recommendations'):
        print('  ERROR: 未找到适用的工艺路线')
        return {}

    top_route = routes['recommendations'][0]
    print(f'  推荐工艺: {top_route["route_name_zh"]} (评分: {top_route["total_score"]})')
    for reason in top_route.get('suitability_reasons', [])[:3]:
        print(f'    + {reason}')

    # ---- Step 2: Design Calculation ----
    print('[2/5] 设计计算...')
    target_limits = kb.resolve_standard_limits(proj['target_standard'])
    orchestrator = CalculationOrchestrator(kb, registry)

    calc_inputs = [{
        'unit_code': u['unit_code'],
        'unit_name_zh': u['unit_name_zh'],
        'sequence_order': u.get('sequence', i+1),
        'is_mandatory': u.get('is_mandatory', True),
    } for i, u in enumerate(top_route.get('units', []))]

    calc_results = orchestrator.run_route(
        route_units=calc_inputs,
        raw_water=wq,
        flow_rate=flow_m3_d,
        flow_rate_peak=flow_m3_d * 1.3,
        target_standard_limits=target_limits,
        design_temp=proj.get('design_temp', 20),
    )
    print(f'  完成 {len(calc_results)} 个构筑物计算')
    for r in calc_results:
        warnings = len(r.warnings) if r.warnings else 0
        flag = ' [!]' if warnings > 0 else ''
        print(f'    {r.unit_name_zh}: {len(r.computed_params)} params, {warnings} warnings{flag}')

    # ---- Step 3: Drawing Analysis (optional) ----
    drawing_elements = []
    if opts.get('include_drawing_analysis', True) and drawings_cfg:
        print('[3/5] 图纸解析...')
        drawing_elements = parse_drawings(drawings_cfg)

    # ---- Step 4: Equipment Selection (optional) ----
    equipment_list = []
    if opts.get('include_equipment_selection', True):
        print('[4/5] 设备选型...')
        eq_selector = EquipmentSelector(kb)
        calc_data = [{
            'calculator_code': r.unit_code,
            'output_parameters': r.computed_params,
        } for r in calc_results]

        route_units_data = [{
            'unit_code': u['unit_code'],
            'unit_name_zh': u['unit_name_zh'],
        } for u in top_route.get('units', [])]

        eq_result = eq_selector.select(
            route_units=route_units_data,
            calculation_results=calc_data,
            flow_rate=flow_m3_d,
        )
        equipment_list = eq_result.get('equipment_list', [])
        total_eq = eq_result.get('summary', {}).get('total_equipment_cost', 0)
        print(f'  已选 {len(equipment_list)} 项设备, 总价约 {total_eq:,.0f} 元')

        # 自动校核
        verifier = EquipmentVerifier()
        verify = verifier.verify(equipment_list, calc_data)
        vsum = verify['summary']
        print(f'  校核: {vsum.get("pass",0)}通过 {vsum.get("warn",0)}警告 {vsum.get("fail",0)}不通过')
        for item in verify['items']:
            if item['overall'] != 'pass':
                print(f'    [{"FAIL" if item["overall"]=="fail" else "WARN"}] {item["model_name"][:30]}')

    # ---- Step 5: Cost Estimation (optional) ----
    cost_data = None
    if opts.get('include_cost_estimation', True):
        print('[5/5] 造价估算...')
        estimator = CostEstimator(kb)
        calc_data_for_cost = [{
            'calculator_code': r.unit_code,
            'output_parameters': r.computed_params,
        } for r in calc_results]

        cost_data = estimator.estimate(
            flow_rate=flow_m3_d,
            calculation_results=calc_data_for_cost,
            equipment_list=equipment_list if equipment_list else [],
        )
        capex = cost_data.get('capex', {})
        opex = cost_data.get('opex', {})
        print(f'  CAPEX: {capex.get("total_capex",0):,.2f} 万元')
        print(f'  OPEX:  {opex.get("total_annual_opex",0):,.2f} 万元/年')
        print(f'  吨水成本: {cost_data.get("cost_per_m3",0):.2f} 元/m3')

    # ---- Parse equipment list from Excel ----
    equip_from_xlsx = []
    if equip_path:
        equip_from_xlsx = parse_equipment_list(equip_path)

    return {
        'project': proj,
        'water_quality': wq,
        'flow_m3_d': flow_m3_d,
        'target_limits': target_limits,
        'selected_route': top_route,
        'all_routes': routes,
        'calculation_results': calc_results,
        'drawing_elements': drawing_elements,
        'equipment_list': equipment_list,
        'equipment_xlsx': equip_from_xlsx,
        'cost_data': cost_data,
    }


def generate_docx(results: dict, output_path: str = None):
    """Generate DOCX using the standard template format."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    proj = results['project']
    wq = results['water_quality']
    route = results['selected_route']
    calcs = results['calculation_results']
    cost = results.get('cost_data')
    equip = results.get('equipment_list', [])
    flow = results['flow_m3_d']

    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21); s.page_height = Cm(29.7)
        s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.paragraph_format.line_spacing = 1.5

    def R(p, text, bold=False, sz=Pt(11), fn='SimSun', align=None, color=None):
        p.alignment = align
        r = p.add_run(text); r.bold = bold; r.font.size = sz
        r.font.name = fn; r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        if color: r.font.color.rgb = color
        return r

    def H1(text):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(24)
        R(p, text, True, Pt(15), 'SimHei', WD_ALIGN_PARAGRAPH.CENTER)

    def H2(text):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18)
        R(p, text, True, Pt(13), 'SimHei')

    def H3(text):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12)
        R(p, text, True, Pt(12), 'SimHei')

    def P(text, indent=True):
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        if indent: p.paragraph_format.first_line_indent = Cm(0.74)
        R(p, text, False, Pt(11))

    def I(text):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        R(p, text, False, Pt(11))

    def TBL(caption, headers, rows):
        if caption:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            R(p, caption, True, Pt(10), 'SimHei', WD_ALIGN_PARAGRAPH.CENTER,
              RGBColor(0x1F, 0x4E, 0x79))
        tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
        tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]; c.text = ''
            pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(str(h)); rr.bold = True; rr.font.size = Pt(9)
            rr.font.name = 'SimSun'; rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D9D9D9')
            c._tc.get_or_add_tcPr().append(shd)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = tbl.rows[ri+1].cells[ci]; c.text = ''
                pp = c.paragraphs[0]
                rr = pp.add_run(str(val) if val is not None else '')
                rr.font.size = Pt(9); rr.font.name = 'SimSun'
                rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        doc.add_paragraph()

    def PAGEBR():
        doc.add_page_break()

    # ===== COVER =====
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, proj.get('location', ''), True, Pt(22), 'SimHei')
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, proj['name'], True, Pt(22), 'SimHei')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, '设计方案', True, Pt(16), 'SimHei')
    for _ in range(4): doc.add_paragraph()
    info = [
        f'设计规模：{proj["flow_rate_m3_h"]} m3/h（约 {flow} m3/d）',
        f'废水类型：{proj["wastewater_type"]}',
        f'排放标准：{proj.get("target_standard", "")}',
        f'主体工艺：{route.get("route_name_zh", "")}',
        f'编制日期：{datetime.date.today().strftime("%Y年%m月")}',
    ]
    for line in info:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, line, False, Pt(11))
    PAGEBR()

    # ===== Ch1 =====
    H1('第一章  项目概况')
    H2('1.1  项目概况')
    P(f'本项目位于{proj.get("location","")}，设计处理能力为{proj["flow_rate_m3_h"]} m3/h（约{flow} m3/d），设计24小时连续运行。出水水质执行相关排放标准。')
    H2('1.2  设计原则')
    for i, pr in enumerate([
        '贯彻执行国家及地方关于环境保护的法律法规、规范和标准。',
        '遵循"分类收集、分质处理"的技术路线，确保出水稳定达标。',
        '选择技术先进成熟、运行稳定可靠、自动化程度高的处理工艺。',
        '充分考虑废水特殊性，在设备选材和防腐方面预留安全余量。',
        '优化总体布置，降低工程投资和运行成本。',
        '合理处置污泥等废物，确保环境安全。',
    ], 1): I(f'（{i}）{pr}')
    H2('1.3  设计范围')
    for i, sc in enumerate([
        '从废水处理设施进水口起至标准化排放口止的全部处理构筑物、设备、管道及电气自控系统。',
        '废水处理工程的工艺流程设计、工艺设备选型、构筑物结构设计、电气控制设计。',
        '废水处理设施的设备制造、安装、调试及操作人员培训。',
        '废水处理站的动力配电由业主将主电源引至站区配电控制柜。',
        '本设计不包括厂区废水收集管网及排出界区的外排水管网。',
    ], 1): I(f'（{i}）{sc}')

    # ===== Ch2 =====
    PAGEBR()
    H1('第二章  设计基础资料')
    H2('2.1  废水产生情况')
    P(f'根据企业提供资料，设计处理能力为{proj["flow_rate_m3_h"]} m3/h，设计每天运行时间24 h。')
    TBL('表2-1  设计进水水质',
        ['控制项目', '进水浓度', '单位'],
        [[k, v, 'mg/L' if k not in ('ph','temperature') else ''] for k, v in wq.items()])

    H2('2.2  处理要求')
    P(f'出水执行 {proj.get("target_standard","相关标准")}。')
    limits = results.get('target_limits', {})
    if limits:
        TBL('表2-2  排放标准限值',
            ['控制项目', '排放限值', '单位'],
            [[k, v, 'mg/L' if k not in ('ph','temperature') else ''] for k, v in limits.items()])

    # ===== Ch3 =====
    PAGEBR()
    H1('第三章  处理工艺选择')
    H2('3.1  工艺选择')
    P(f'经技术经济比较，确定采用"{route.get("route_name_zh","")}"处理工艺。')
    for reason in route.get('suitability_reasons', []):
        I(f'+ {reason}')

    H2('3.2  工艺流程')
    # Embed flow diagram
    flow_png = Path(__file__).parent / 'reports' / 'flow_diagram.png'
    if flow_png.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(p, '图3-1  废水处理工艺流程图', True, Pt(10), 'SimHei')
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(flow_png), width=Inches(5.5))

    H2('3.3  工艺流程简述')
    units_desc = route.get('units', [])
    for i, u in enumerate(units_desc, 1):
        P(f'（{i}）{u.get("unit_name_zh", u.get("unit_code",""))}：{u.get("purpose_zh","处理单元")}。')

    H2('3.4  预计处理效率')
    if calcs:
        param_keys = list(calcs[0].computed_params.keys())[:5] if calcs[0].computed_params else []
        rows = []
        for r in calcs:
            row = [r.unit_name_zh]
            for k in param_keys:
                v = r.computed_params.get(k, '-')
                row.append(f'{v:.1f}' if isinstance(v, float) else str(v))
            rows.append(row)
        TBL('表3-1  主要构筑物设计参数',
            ['构筑物'] + param_keys, rows)

    # ===== Ch4 =====
    PAGEBR()
    H1('第四章  构筑物及设备')
    for i, r in enumerate(calcs, 1):
        H3(f'（{i}）{r.unit_name_zh}')
        params = r.computed_params
        if params:
            for pn, pv in params.items():
                I(f'{pn}：{pv:.2f}' if isinstance(pv, float) else f'{pn}：{pv}')

    # ===== Ch5 =====
    PAGEBR()
    H1('第五章  施工计划')
    H2('5.1  施工计划')
    for i, phase in enumerate([
        '施工准备阶段（第1~15天）：人员及机械进场，场地平整，技术交底。',
        '土建施工阶段（第16~90天）：调节池、设备基础等同步施工。',
        '设备安装阶段（第91~135天）：设备安装，管道连接，电气敷设。',
        '调试阶段（第136~165天）：单机试车、联动试车、工艺调试。',
        '验收阶段（第166~180天）：稳定达标运行，竣工验收。',
    ], 1): I(f'（{i}）{phase}')

    # ===== Ch6 =====
    PAGEBR()
    H1('第六章  工程投资估算与运行成本')
    H2('6.1  投资估算')

    total_equip_cost = sum(e.get('total_price_cny', e.get('unit_price_cny', 0) * e.get('quantity', 1)) for e in equip)
    equip_wan = round(total_equip_cost / 10000, 2)

    equip_rows = []
    for i, e in enumerate(equip[:30], 1):
        equip_rows.append([str(i), e.get('model_name_zh', ''), e.get('model_id', ''),
                          str(e.get('quantity', 1)), f'{e.get("unit_price_cny",0):,.0f}',
                          f'{e.get("total_price_cny",0):,.0f}'])
    if equip_rows:
        TBL('表6-1  主要设备清单',
            ['序号', '设备名称', '型号', '数量', '单价(元)', '总价(元)'], equip_rows)
    P(f'设备总价约 {total_equip_cost:,.0f} 元（{equip_wan} 万元）。', indent=False)

    H2('6.2  运行成本')
    if cost:
        capex = cost.get('capex', {})
        opex = cost.get('opex', {})
        TBL('表6-2  运行成本汇总',
            ['成本项目', '年费用（万元）'],
            [['电费', f'{opex.get("energy_cost",0):,.2f}'],
             ['药剂费', f'{opex.get("chemical_cost",0):,.2f}'],
             ['人工费', f'{opex.get("labor_cost",0):,.2f}'],
             ['维护费', f'{opex.get("maintenance_cost",0):,.2f}'],
             ['污泥处置费', f'{opex.get("sludge_disposal_cost",0):,.2f}'],
             ['折旧', f'{opex.get("depreciation_cost",0):,.2f}'],
             ['合计', f'{opex.get("total_annual_opex",0):,.2f}']])
        P(f'吨水运行成本：{cost.get("cost_per_m3",0):.2f} 元/m3。', indent=False)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p, '—— 方案完 ——', False, Pt(10), 'SimSun', color=RGBColor(0x99, 0x99, 0x99))

    if not output_path:
        out_dir = Path(__file__).parent / 'reports'
        out_dir.mkdir(exist_ok=True)
        safe_name = proj['name'].replace(' ','_').replace('/','_')[:30]
        output_path = str(out_dir / f'{safe_name}_{datetime.date.today().strftime("%Y%m%d")}.docx')

    doc.save(output_path)
    print(f'\n  DOCX saved: {output_path}')
    return output_path


# ==================== MAIN ====================
if __name__ == '__main__':
    yaml_path = sys.argv[1] if len(sys.argv) > 1 else 'project_input.yaml'
    if not os.path.exists(yaml_path):
        print(f'ERROR: 找不到配置文件 {yaml_path}')
        print('用法: python generate_report.py [project_input.yaml]')
        sys.exit(1)

    print(f'加载配置: {yaml_path}')
    config = load_input(yaml_path)

    results = run_pipeline(config)
    if not results:
        print('\nERROR: 流水线执行失败')
        sys.exit(1)

    output = config.get('options', {}).get('output_path', '')
    generate_docx(results, output if output else None)
    print('\n[OK] Report generation complete')
