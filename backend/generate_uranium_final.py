"""Generate fully formatted professional uranium wastewater report."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import datetime

doc = Document()

# ---- Page Setup ----
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True

# ---- Styles ----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
style.paragraph_format.line_spacing = 1.5

# ---- Color Scheme ----
HEADER_BG = '1F4E79'       # Dark blue for main headers
HEADER_BG2 = '2E75B6'      # Medium blue for sub headers
HEADER_BG3 = 'D6E4F0'      # Light blue for table headers
ACCENT = '1F4E79'
BORDER_COLOR = 'B4C6E7'
ROW_ALT = 'F2F7FC'         # Alternating row

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement('w:{}'.format(edge))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', BORDER_COLOR))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, size=Pt(11), font_cn='SimSun', font_en='Times New Roman', color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    if color:
        run.font.color.rgb = color
    return run

def add_h1(text):
    """Chapter heading - large bold centered"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(16)
    add_run(p, text, bold=True, size=Pt(16), font_cn='SimHei', color=RGBColor(0x1F, 0x4E, 0x79))
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(text):
    """Section heading"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, bold=True, size=Pt(13), font_cn='SimHei', color=RGBColor(0x2E, 0x75, 0xB6))
    return p

def add_h3(text):
    """Sub heading"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size=Pt(12), font_cn='SimHei', color=RGBColor(0x33, 0x33, 0x33))
    return p

def add_p(text, indent=True):
    """Body paragraph"""
    p = doc.add_paragraph()
    add_run(p, text, size=Pt(10.5))
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.6
    return p

def add_item(text):
    """Indented item with numbering"""
    p = doc.add_paragraph()
    add_run(p, text, size=Pt(10.5))
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_caption(text):
    """Table/figure caption"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size=Pt(10), font_cn='SimHei', color=RGBColor(0x1F, 0x4E, 0x79))
    return p

def add_tbl(caption, headers, rows):
    """Professional formatted table with caption, shaded header, alternating rows"""
    if caption:
        add_caption(caption)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(h), bold=True, size=Pt(9), font_cn='SimHei', color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            # Center-align numeric columns, left-align text
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif val and str(val).replace(',','').replace('.','').replace('-','').isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, str(val) if val is not None else '', size=Pt(9))
            # Alternating row colors
            if ri % 2 == 1:
                set_cell_shading(cell, ROW_ALT)
            # Bold total rows
            if str(val) in ('小计', '合  计', '合计') or (ri > 0 and rows[ri-1][0] == '' and str(val) == ''):
                for run in p.runs:
                    run.bold = True

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return table

def page_break():
    doc.add_page_break()

# ==================== COVER PAGE ====================
# Top decorative bar
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '36')
bottom.set(qn('w:color'), '1F4E79')
pBdr.append(bottom)
pPr.append(pBdr)

for _ in range(4):
    doc.add_paragraph()

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '内蒙古包头市', bold=True, size=Pt(26), font_cn='SimHei', color=RGBColor(0x1F, 0x4E, 0x79))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '含铀废水处理工程', bold=True, size=Pt(26), font_cn='SimHei', color=RGBColor(0x1F, 0x4E, 0x79))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '设计方案', bold=True, size=Pt(20), font_cn='SimHei', color=RGBColor(0x2E, 0x75, 0xB6))

for _ in range(5):
    doc.add_paragraph()

# Info table on cover
info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('工程规模', '12 m3/h（约288 m3/d）'),
    ('废水类型', '含铀工业废水'),
    ('排放标准', 'GB 8978-1996 一级 + GB 23727-2009'),
    ('主体工艺', '化学沉淀+DTRO/RO+铵盐沉淀+硅胶吸附+MVR+生化'),
    ('编制单位', ''),
    ('编制日期', datetime.date.today().strftime('%Y年%m月')),
]
for i, (label, value) in enumerate(info_data):
    c0 = info_table.cell(i, 0)
    c0.text = ''
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, label + '：', bold=True, size=Pt(11), font_cn='SimHei')
    c0.width = Cm(3.5)

    c1 = info_table.cell(i, 1)
    c1.text = ''
    p = c1.paragraphs[0]
    add_run(p, value, size=Pt(11))
    c1.width = Cm(7)

# Remove table borders for cover
for row in info_table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            e = OxmlElement('w:{}'.format(edge))
            e.set(qn('w:val'), 'none')
            e.set(qn('w:sz'), '0')
            e.set(qn('w:color'), 'auto')
            tcBorders.append(e)
        tcPr.append(tcBorders)

page_break()

# ==================== TOC ====================
add_h1('目    录')

toc_data = [
    ('第一章  项目概况', [
        '1.1  项目概况',
        '1.2  设计原则',
        '1.3  设计范围',
    ]),
    ('第二章  设计基础资料', [
        '2.1  废水产生情况',
        '2.2  处理要求',
    ]),
    ('第三章  处理工艺选择', [
        '3.1  废水特点及处理难点',
        '3.2  处理工艺流程',
        '3.3  工艺流程简述',
        '3.4  预计处理效率',
        '3.5  平面布置',
    ]),
    ('第四章  构筑物及设备', [
        '4.1  预处理段',
        '4.2  化学沉淀段',
        '4.3  膜浓缩段（DTRO+RO）',
        '4.4  浓水除铀段（铵盐沉淀+硅胶吸附）',
        '4.5  除氟脱氨蒸发段',
        '4.6  生化处理段',
        '4.7  污泥处理段',
    ]),
    ('第五章  施工计划与组织设计', [
        '5.1  施工计划',
        '5.2  施工组织设计',
    ]),
    ('第六章  工程投资估算与运行成本', [
        '6.1  估算依据',
        '6.2  投资估算',
        '6.3  运行成本',
    ]),
]

for chapter, sections in toc_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    add_run(p, chapter, bold=True, size=Pt(12), font_cn='SimHei')
    for sec in sections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        add_run(p, sec, size=Pt(11))

page_break()

# ==================== CHAPTER 1 ====================
add_h1('第一章  项目概况')

add_h2('1.1  项目概况')
add_p('本项目位于内蒙古自治区包头市，为含铀工业废水处理工程。包头市是我国重要的稀土及核工业基地，含铀废水主要来源于铀矿冶、稀土冶炼分离及核燃料加工等工艺环节。废水中铀主要以六价铀酰离子（UO2(2+)）形态存在，同时含有氟化物、氨氮、硝酸盐、有机物等多种特征污染物。')
add_p('项目设计处理能力为12 m3/h（约288 m3/d），设计24小时连续运行。出水水质执行《污水综合排放标准》（GB 8978-1996）表4中一级标准，同时满足《铀矿冶辐射防护和环境保护规定》（GB 23727-2009）中铀的排放限值要求（总铀≤0.05 mg/L）。处理后达标废水排入园区污水管网。')

add_h2('1.2  设计原则')
for i, pr in enumerate([
    '贯彻执行国家及地方关于环境保护和放射性废物管理的法律法规、规范和标准。',
    '遵循"分类收集、分质处理、浓淡分离、深度净化"的技术路线，确保出水稳定达标。',
    '选择技术先进成熟、运行稳定可靠、自动化程度高的处理工艺，减少人为操作误差。',
    '充分考虑含铀废水的特殊性（放射性、腐蚀性、高盐分），在设备选材和防腐方面预留足够安全余量。',
    '优化总体布置，减少土建工程量，降低工程投资和运行成本。',
    '合理处置含铀污泥等放射性废物，确保环境安全。',
], 1):
    add_item('（%d）%s' % (i, pr))

add_h2('1.3  设计范围')
for i, sc in enumerate([
    '从废水处理设施进水口起至标准化排放口止的全部处理构筑物、设备、管道及电气自控系统。',
    '废水处理工程的工艺流程设计、工艺设备选型、构筑物结构设计、电气控制设计。',
    '废水处理设施的设备制造、安装、调试及操作人员培训。',
    '废水处理站的动力配电由业主将主电源引至站区配电控制柜，柜后配电由承包方负责。',
    '本设计不包括厂区废水收集管网及排出界区的外排水管网。',
    '本设计不包括含铀污泥的最终处置（需委托有资质的放射性废物处置单位）。',
], 1):
    add_item('（%d）%s' % (i, sc))

# ==================== CHAPTER 2 ====================
page_break()
add_h1('第二章  设计基础资料')

add_h2('2.1  废水产生情况')
add_p('根据企业提供资料及同类含铀废水项目经验，含铀废水主要来源于铀矿冶浸出液、稀土冶炼分离萃余液、核燃料元件加工清洗水等工艺环节。废水排放方式为连续排放，各工艺段废水水质水量有所波动，需经调节池均质均量后进入处理系统。设计处理能力为12 m3/h，设计每天运行时间24 h。')

add_tbl('表2-1  设计进水水质',
    ['控制项目', '进水浓度范围', '备注'],
    [['pH', '3 ~ 9', '视工艺段波动'],
     ['CODcr（mg/L）', '500 ~ 2000', '含有机物及萃余液'],
     ['NH3-N（mg/L）', '100 ~ 500', '铵盐工艺带入'],
     ['总氮（mg/L）', '150 ~ 800', '硝酸盐 + 氨氮'],
     ['SS（mg/L）', '200 ~ 500', '悬浮物'],
     ['氟化物（mg/L）', '20 ~ 100', '含氟工艺废水'],
     ['总铀 U（mg/L）', '5 ~ 50', '核心去除目标'],
     ['总a放射性（Bq/L）', '10 ~ 100', '放射性指标'],
     ['总b放射性（Bq/L）', '10 ~ 100', '放射性指标'],
     ['Ca2+/Mg2+（mg/L）', '200 ~ 500', '硬度离子'],
     ['TDS（mg/L）', '3000 ~ 15000', '高盐分']])

add_h2('2.2  处理要求')
add_p('本项目出水执行《污水综合排放标准》（GB 8978-1996）表4中一级标准，其中总铀指标执行《铀矿冶辐射防护和环境保护规定》（GB 23727-2009）的排放限值。')

add_tbl('表2-2  排放标准限值',
    ['控制项目', '排放限值', '依据标准'],
    [['pH', '6 ~ 9', 'GB 8978-1996 一级'],
     ['CODcr（mg/L）', '≤ 100', 'GB 8978-1996 一级'],
     ['BOD5（mg/L）', '≤ 30', 'GB 8978-1996 一级'],
     ['SS（mg/L）', '≤ 70', 'GB 8978-1996 一级'],
     ['NH3-N（mg/L）', '≤ 15', 'GB 8978-1996 一级'],
     ['总氮（mg/L）', '≤ 45（参考）', '地方环保要求'],
     ['氟化物（mg/L）', '≤ 10', 'GB 8978-1996 一级'],
     ['总铀 U（mg/L）', '≤ 0.05', 'GB 23727-2009'],
     ['总a放射性（Bq/L）', '≤ 1', 'GB 8978-1996 一级'],
     ['总b放射性（Bq/L）', '≤ 10', 'GB 8978-1996 一级']])

# ==================== CHAPTER 3 ====================
page_break()
add_h1('第三章  处理工艺选择')

add_h2('3.1  废水特点及处理难点')
add_p('含铀废水属于特种工业废水，具有以下特点和难点：')
for i, pt in enumerate([
    '铀的深度去除难度极大：进水铀浓度5~50 mg/L，要求出水≤0.05 mg/L，去除率需达99%以上，单一处理技术无法实现。',
    '高盐分干扰生化处理：废水TDS高达3000~15000 mg/L，常规活性污泥法无法耐受，需通过膜分离和蒸发实现脱盐后再进入生化系统。',
    '氟离子腐蚀性强：F-浓度20~100 mg/L，对碳钢和普通不锈钢均有腐蚀，需采用玻璃钢、ECTFE喷涂、2205双相不锈钢等特种防腐材料。',
    '放射性废物管理要求严格：含铀污泥和浓缩液属于放射性废物，需按GB 14500-2002要求进行专门管理和处置。',
    '污染物种类多、协同去除难度大：需同时去除铀、COD、NH3-N、TN、F-、SS等多种污染物，各处理单元需优化组合。',
], 1):
    add_item('（%d）%s' % (i, pt))

add_h2('3.2  处理工艺流程')
add_p('根据含铀废水的特点及处理要求，经技术经济比较，确定采用"化学混凝沉淀+多介质过滤+DTRO+RO双膜浓缩+浓水铵盐沉淀+硅胶吸附除铀+产水深度处理（除氟除硬+脱氨+MVR蒸发+水解酸化+A/O+MBR+芬顿氧化）"的全链条处理工艺。')

# ---- ProcessOn-style Flow Diagram (embedded PNG image) ----
add_caption('图3-1  废水处理工艺流程图')

# Embed the ProcessOn-style flow diagram image
flow_img_path = r'E:\claude\wastewater-app\backend\reports\flow_diagram.png'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(flow_img_path, width=Inches(6.1))

add_p('')

add_h2('3.3  工艺流程简述')
for i, desc in enumerate([
    '废水经收集管道进入原水调节池（有效容积约100 m3，HRT=8 h），池内设穿孔曝气搅拌装置，对来水进行均质均量调节。调节池出水通过提升泵送至综合调节池。',
    '综合调节池（有效容积约150 m3，HRT=12.5 h）进一步均化水质，同时接纳MVR蒸发冷凝水等回流液。池内设穿孔曝气搅拌系统，必要时投加NaOH调节pH。',
    '综合调节池出水通过提升泵进入絮凝沉淀池。向池内投加PAC（聚合氯化铝）和PAM（聚丙烯酰胺），同时投加CaO调节pH并部分去除氟离子（生成CaF2沉淀）。沉淀池有效容积48 m3，表面负荷0.67 m3/(m2-h)，HRT=4 h。',
    '絮凝沉淀池出水自流进入中间水箱，经提升泵加压后进入多介质过滤器（o1400x2850，2台），滤速8 m/h，进一步去除残余悬浮物，降低后续膜系统的SDI值。过滤器定期气水反冲洗。',
    '过滤出水进入DTRO（碟管式反渗透）装置。在75 bar操作压力下运行，回收率75%。DTRO膜对铀的截留率>99%，对盐分的截留率>95%。产水进入RO装置，浓水进入铵盐沉淀段。',
    'DTRO产水进入RO（反渗透）装置。一级二段式设计，回收率85%，操作压力15~20 bar。RO产水达排放标准，RO浓水返回铵盐沉淀段。双膜系统总回收率约94%。',
    'DTRO+RO浓水（约2.6 m3/h）进入铵盐沉淀反应槽。酸性条件下加氨水和硝酸，铀酰离子与铵根反应生成重铀酸铵((NH4)2U2O7)黄色沉淀。板框压滤（20 m2）固液分离，滤饼回收铀。',
    '压滤滤液经10 um和5 um袋式过滤后进入硅胶吸附柱。8台吸附柱（1.5 m3/柱）串联，pH 5~6条件下硅胶高效吸附U(VI)，尾液铀浓度降至0.01 mg/L以下。饱和硅胶稀硝酸解吸再生。',
    'RO产水（约7.65 m3/h）与吸附尾液合并，进入除氟除硬沉淀池。投加Ca(OH)2和Na2CO3去除F-（CaF2）和钙镁硬度（CaCO3/Mg(OH)2）。两级沉淀池（单池9.0x2.5x3.0 m），HRT=3 h。',
    '除硬后废水进入脱氨塔（6 m3/h），蒸汽汽提将游离氨转移至气相，稀硫酸吸收生成硫酸铵。脱氨后NH3-N降至50 mg/L以下。',
    '脱氨废水进入MVR蒸发装置（5.4 m3/h，2205双相不锈钢），真空低温蒸发。冷凝水（约5 m3/h）自流至综合调节池；浓缩液（0.45 m3/h）进入刮板蒸发干化，产出结晶盐外运。',
    'MVR冷凝水与原水混合后（TDS<3000 mg/L）进入生化系统：水解酸化池（HRT=11 h）分解大分子；缺氧池（HRT=11 h）反硝化脱氮；好氧池（HRT=22 h）硝化+有机物降解；MBR膜池（PTFE 960 m2）泥水分离，MLSS 8000~12000 mg/L。',
    'MBR出水进入芬顿氧化塔（HRT=3.5 h），酸性条件（pH 3~4）下H2O2+FeSO4产生-OH自由基氧化残余COD至100 mg/L以下。回调pH至中性，铁盐絮体沉淀去除。出水达标排放。',
    '各沉淀池排泥、MBR剩余污泥及芬顿化学污泥排入污泥池（8.0x3.0x4.0 m），浓缩后柱塞泵送入板框压滤机（80 m2隔膜式）压滤。含铀泥饼作为放射性废物，专用暂存间存放，定期委托有资质单位处置。滤液回流至综合调节池。',
], 1):
    add_p('（%d）%s' % (i, desc))

add_h2('3.4  预计处理效率')
add_tbl('表3-1  主要污染物分级处理效率表（单位：mg/L，铀除外）',
    ['处理单元', 'COD', 'NH3-N', 'TN', 'SS', 'F-', '总铀 U*'],
    [['进水', '500-2000', '100-500', '150-800', '200-500', '20-100', '5-50'],
     ['絮凝沉淀', '200-500', '80-250', '120-600', '20-50', '5-15', '1-10'],
     ['多介质过滤', '150-350', '80-250', '120-600', '5-10', '3-10', '0.5-5'],
     ['DTRO', '10-30', '2-8', '3-15', '<1', '<1', '0.01-0.05'],
     ['RO', '<10', '<2', '<5', '<1', '<1', '<0.01'],
     ['硅胶吸附**', '—', '—', '—', '—', '—', '<0.01'],
     ['除氟除硬', '<10', '<2', '<5', '<1', '<1', '<0.01'],
     ['脱氨塔', '<10', '<30', '<30', '<1', '<1', '<0.01'],
     ['MVR冷凝水', '<20', '<30', '<30', '<1', '<1', '<0.01'],
     ['生化+MBR', '<60', '<8', '<20', '<5', '<1', '<0.01'],
     ['芬顿氧化', '<80', '<8', '<20', '<10', '<1', '<0.01'],
     ['排放限值', '100', '15', '45', '70', '10', '0.05']])
add_p('注：* 铀浓度单位为 mg/L；** 硅胶吸附处理对象为DTRO/RO浓水，上表为吸附后尾液合并至主流程后的折算值。')

add_h2('3.5  平面布置')
add_p('废水处理站根据功能分为以下区域：')
for i, area in enumerate([
    '预处理区：布置原水调节池（8.0x4.0x4.0 m地下钢砼）和综合调节池（8.0x6.0x4.0 m地下钢砼）。',
    '化学处理及膜车间：布置絮凝沉淀池（9.0x3.0x3.5 m）、多介质过滤器（o1400x2850，2台）、DTRO（2套）、RO（2套）及相关泵组和加药装置。设于操作用房内（9.0x4.0x3.0 m）。',
    '浓水处理区：布置铵盐沉淀反应槽（4台）、板框压滤机、袋式过滤器、硅胶吸附柱（8台）及相关储槽和泵组。',
    '生化处理区：布置水解酸化池、缺氧池、好氧池（各10.0x3.0x5.0 m，碳钢FRP防腐）、MBR膜池、芬顿氧化塔及鼓风机房。',
    '蒸发结晶区：布置脱氨塔、MVR蒸发装置、刮板蒸发干化装置及氨水暂存罐等。',
    '污泥处理区：布置污泥池（8.0x3.0x4.0 m）和板框压滤机（2台，80 m2）。',
    '附属设施：操作用房、控制室、配电间、在线监测室。',
], 1):
    add_item('（%d）%s' % (i, area))

# ==================== CHAPTER 4 ====================
page_break()
add_h1('第四章  构筑物及设备')

def add_unit(title, count, structure, size, params, equipment):
    add_h3(title)
    add_item('数量：' + count)
    add_item('结构形式：' + structure)
    add_item('池体尺寸：' + size)
    if params and params != '—':
        add_item('设计参数：' + params)
    if equipment:
        for i, eq in enumerate(equipment, 1):
            add_item('配套设备%d：%s' % (i, eq))

add_h2('4.1  预处理段')
add_unit('（1）原水调节池',
    count='1座', structure='地下式钢砼防腐',
    size='8.0 m x 4.0 m x 4.0 m（有效水深3.5 m，有效容积约100 m3）',
    params='停留时间8 h',
    equipment=[
        '提升水泵2台。Q=12 m3/h，H=15 m，N=1.1 kW，自吸泵，过流304不锈钢。',
        '穿孔曝气搅拌系统1套。DN80，ABS材质。',
        '液位控制系统1套（浮球液位计，0~4 m）。'])

add_unit('（2）综合调节池',
    count='1座', structure='地下式钢砼防腐',
    size='8.0 m x 6.0 m x 4.0 m（有效水深3.5 m，有效容积约150 m3）',
    params='停留时间12.5 h',
    equipment=[
        '废水提升泵2台。Q=12 m3/h，H=15 m，N=1.1 kW，自吸泵，过流304不锈钢。',
        '穿孔曝气搅拌系统1套。DN80，ABS材质，曝气强度0.5~1.0 m3/(m2-h)。',
        '液位控制系统1套（浮球液位计）。'])

add_h2('4.2  化学沉淀段')
add_unit('（3）絮凝沉淀池',
    count='1座（含反应区+斜管沉淀区）',
    structure='地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='9.0 m x 3.0 m x 3.5 m（有效容积48 m3）',
    params='表面负荷0.67 m3/(m2-h)，停留时间4 h',
    equipment=[
        '反应搅拌机3台。ZJ270，轴和叶片衬塑防腐，N=1.1 kW。',
        '斜管填料2套。DN80 PP斜管，60度安装，碳钢防腐支架。',
        '中心筒2套。o300x2000，碳钢防腐。',
        '出水堰1套。8 mm锯齿三角堰，碳钢玻璃钢防腐。',
        '排泥泵2台。Q=10 m3/h，H=25 m，N=1.1 kW，过流304。',
        '加药装置3套。V=1.5 m3，Q=140 L/h，P=0.6 MPa，一箱二泵。分别投加PAC、PAM、CaO/NaOH。',
        'pH在线监测仪1套。'])

add_h2('4.3  膜浓缩段')
add_unit('（4）多介质过滤系统',
    count='1套（含2台过滤器）', structure='碳钢衬胶防腐（3+2 mm）',
    size='o1400 x 2850 mm（单台）',
    params='滤速8 m/h，反冲洗强度10 L/(m2-s)',
    equipment=[
        '多介质过滤器2台。o1400x2850，含石英砂+无烟煤滤料、自动阀门。',
        '提升水泵2台。Q=12 m3/h，H=25 m，N=2.2 kW，离心泵，过流304。',
        '反冲洗水泵2台。Q=60 m3/h，H=20 m，N=5.5 kW，离心泵，过流304。',
        '中间水箱1台。10 m3，PE储罐。'])

add_unit('（5）DTRO装置',
    count='2套（并联）', structure='支架304不锈钢，碟管式反渗透膜组件',
    size='—', params='单套12 m3/h，回收率75%，操作压力75 bar',
    equipment=[
        'DTRO膜装置2套。含膜柱、高压泵（75 bar）、保安过滤器（5 um）、膜清洗装置、在线仪表。',
        'DTRO进水泵2台。Q=12 m3/h，H=25 m，N=2.2 kW，过流304。',
        'DTRO产水水箱1台。10 m3，PE储罐。'])

add_unit('（6）RO装置',
    count='2套（并联）', structure='支架304不锈钢，卷式反渗透膜组件',
    size='—', params='单套9 m3/h，一级二段式，回收率85%',
    equipment=[
        'RO膜装置2套。含膜元件、高压泵、保安过滤器（5 um）、膜清洗装置、在线仪表。',
        'RO进水泵2台。Q=10 m3/h，H=25 m，N=2.2 kW，过流304。',
        '浓水水箱1台。10 m3，PE储罐。'])

add_h2('4.4  浓水除铀段')
add_unit('（7）铵盐沉淀反应系统',
    count='1套', structure='反应槽利旧4台，高位槽Q345R+ECTFE喷涂',
    size='反应槽单台>=6 m3；高位槽各1 m3', params='—',
    equipment=[
        '浓水输送泵2台。Q=6 m3/h，H=20 m，N=1.1 kW，过流304。',
        '氨水高位槽1台。1 m3，Q345R+ECTFE。',
        '浓硝酸高位槽1台。1 m3，Q345R+ECTFE。',
        '板框压滤机1台。过滤面积20 m2。',
        '袋式过滤器2台(10 um)+1台(5 um)。Q=10 m3/h，过流304。',
        '滤液储槽1台。5 m3，Q345R+ECTFE。'])

add_unit('（8）硅胶吸附系统',
    count='1套', structure='硅胶吸附柱8台',
    size='单柱容积约1.5 m3', params='—',
    equipment=[
        '齿轮计量泵4台。Q=0.8 m3/h，H=25 m。',
        '硅胶吸附柱8台。单柱1.5 m3。',
        '吸附尾液储槽4台（利旧）。>=5 m3/台。',
        '稀硝酸储槽1台（利旧）。>=2 m3。',
        '解吸液储槽2台（利旧）。>=2 m3/台。',
        '解吸液输送泵4台。Q>=6.3 m3/h，H>=15 m，过流304。',
        '在线铀浓度监测仪1套（激光荧光法）。',
        '在线氨氮监测仪1套（水杨酸法）。'])

page_break()
add_h2('4.5  除氟脱氨蒸发段')
add_unit('（9）除氟除硬沉淀池',
    count='2套', structure='地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='9.0 m x 2.5 m x 3.0 m（单套）',
    params='表面负荷0.96 m3/(m2-h)，停留时间3 h',
    equipment=[
        '反应搅拌机6台。ZJ270，轴和叶片衬塑防腐。',
        '斜管填料4套。DN80 PP，60度安装。',
        '中心筒4套。o300x2000，碳钢防腐。',
        '出水堰4套。8 mm锯齿三角堰，碳钢FRP防腐。',
        '排泥泵4台。Q=10 m3/h，H=25 m，N=1.1 kW。',
        '加药装置4套。V=1.5 m3，一箱二泵。分别投加Ca(OH)2、Na2CO3、PAC、PAM。',
        '中间水箱1台。10 m3，Q345R+ECTFE。',
        '提升泵2台。Q=6 m3/h，H=18 m，N=0.75 kW。'])

add_unit('（10）脱氨塔',
    count='1套', structure='304不锈钢', size='处理能力6 m3/h', params='—',
    equipment=[
        '脱氨塔进水泵2台。Q=6 m3/h，H=18 m，N=0.75 kW，过流304。',
        '脱氨塔本体1套。含塔体、填料、分布器、冷凝器。',
        '氨水暂存罐1只。V=10 m3，PE。',
        '脱氨废水水槽1台。10 m3，PE。'])

add_unit('（11）MVR蒸发系统',
    count='1套', structure='2205双相不锈钢', size='处理能力5.4 m3/h',
    params='MVR机械蒸汽再压缩蒸发',
    equipment=[
        'MVR蒸发装置1套。Q=5.4 m3/h，含蒸发器本体、蒸汽压缩机、预热器、冷凝器、循环泵、真空系统、循环冷却水系统。',
        '刮板蒸发干化装置1套。Q=0.45 m3/h，含刮板蒸发器、冷凝器、循环冷却水系统。',
        '提升泵2台。Q=6 m3/h，H=18 m，N=0.75 kW。'])

add_h2('4.6  生化处理段')
add_unit('（12）水解酸化池',
    count='1座', structure='地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='10.0 m x 3.0 m x 5.0 m', params='停留时间11 h，有效容积约135 m3',
    equipment=[
        '潜水搅拌机2台。QJB400，N=1.1 kW，304不锈钢。',
        '布水装置1套。DN50穿孔布水管，304不锈钢。',
        '水解池循环泵2台。Q=40 m3/h，H=15 m，N=3.0 kW，管道泵，过流304。'])

add_unit('（13）缺氧池',
    count='1台', structure='地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='10.0 m x 3.0 m x 5.0 m',
    params='停留时间11 h，有效容积约135 m3，MLSS 3000~4000 mg/L',
    equipment=['潜水搅拌机2台。QJB400，N=1.1 kW，304不锈钢。'])

add_unit('（14）好氧池+MBR膜池',
    count='好氧池2台 + MBR膜组件4套',
    structure='好氧池：地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='好氧池单台10.0 m x 3.0 m x 5.0 m（有效水深4.5 m，总有效容积约270 m3）',
    params='停留时间22 h，MLSS 4000~5000（好氧）/8000~12000（MBR），DO 2~4 mg/L，污泥龄25~30 d',
    equipment=[
        '微孔曝气器1套。BNQZ-192，200套，ABS组合件。',
        '硝化液回流泵2台。Q=40 m3/h，H=15 m，N=3.0 kW，管道泵，过流304。',
        '污泥回流泵2台。Q=15 m3/h，H=15 m，N=1.5 kW，管道泵，过流304。',
        'MBR膜组件4套。PTFE中空纤维膜，304支架，单套240 m2，总960 m2，通量12.5 L/(m2-h)。',
        '抽吸泵2台。Q=12 m3/h，H=15 m，N=1.1 kW，自吸泵，过流304。',
        '膜清洗装置1套。V=10 m3，PE，含清洗泵Q=40 m3/h，H=15 m，N=3.0 kW。',
        '鼓风机2台（1用1备）。空气悬浮离心风机，Q=10 m3/min，升压60 kPa，N=11 kW。',
        'DO在线监测仪、pH在线监测仪各1套。',
        '加药装置1套（碱度补充）。投加NaOH/Na2CO3。'])

add_unit('（15）芬顿氧化塔',
    count='2套', structure='地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='9.0 m x 3.0 m x 5.0 m（单套）', params='停留时间3.5 h',
    equipment=[
        '加药装置5套。V=1.5 m3，一箱二泵。分别投加H2O2、FeSO4、H2SO4（调酸）、NaOH（中和）、PAM。',
        'pH在线监测仪2套（反应区+中和区）。',
        'ORP在线监测仪1套。'])

add_h2('4.7  污泥处理段')
add_unit('（16）污泥处理系统',
    count='1套', structure='污泥池：地上式钢结构，8 mm碳钢非标，内壁玻璃钢防腐',
    size='污泥池 8.0 m x 3.0 m x 4.0 m', params='—',
    equipment=[
        '污泥池搅拌机2台。JBJ-1200，轴和叶片衬塑防腐，N=3 kW。',
        '污泥泵2台（柱塞泵）。YB150，Q=10 m3/h，P=2.0 MPa，N=7.5 kW。',
        '板框压滤机2台。80 m2，程控隔膜式，碳钢防腐，操作压力0.8 MPa，压榨1.2 MPa，滤板1250x1250。',
        '污泥泥斗2台。非标配套，碳钢防腐。',
        '水仓1台。V=5 m3，304不锈钢。',
        '高压泵2台。Q=5 m3/h，H=140 m，N=5.5 kW。',
        '压榨水箱1台。304不锈钢。',
        '压榨水泵2台。Q=2 m3/h，H=120 m，N=2.2 kW。'])

# ==================== CHAPTER 5 ====================
page_break()
add_h1('第五章  施工计划与组织设计')

add_h2('5.1  施工计划')
add_p('本工程为新建项目，施工期间企业正常生产。施工总工期计划为180日历天，各阶段计划如下：')
for i, phase in enumerate([
    '施工准备阶段（第1~15天）：施工人员及机械进场，搭建临时设施，场地平整。施工区域围挡，设计技术交底，图纸会审。',
    '土建施工阶段（第16~90天）：调节池、设备基础、操作用房同步施工。钢砼水池养护>=14天。设备基础混凝土强度达设计值75%后方可安装设备。',
    '设备加工及进场阶段（第30~90天）：非标钢制罐体工厂制作+FRP防腐内衬，运至现场吊装就位。DTRO/RO/MVR定型设备按计划采购进场。',
    '设备安装阶段（第91~135天）：各单元设备安装，管道阀门连接，电气线路敷设，PLC控制系统安装接线。',
    '调试阶段（第136~165天）：单机试车、清水联动试车、工艺调试。逐步引入实际废水，接种活性污泥，培养驯化微生物。',
    '验收阶段（第166~180天）：出水连续稳定达标运行>=15天，整理竣工资料，组织竣工验收。',
], 1):
    add_item('（%d）%s' % (i, phase))

add_h2('5.2  施工组织设计')
add_p('本工程施工过程分为三部分：')

add_h3('（一）土建工程')
for i, item in enumerate([
    '污水站土建工程包括新建调节池、设备基础、操作用房等钢砼/砖混构筑物。',
    '施工前对场地进行平整，根据设计图纸进行测量放线。',
    '水池施工按GB 50141-2008执行，确保防水防腐质量。',
    '钢砼水池浇筑完成后应及时养护，养护期内严禁施加荷载。',
], 1):
    add_item('%d）%s' % (i, item))

add_h3('（二）设备与安装工程')
for i, item in enumerate([
    '设备与安装工程包含废水处理系统的全部工艺设备、管道阀门、电气自控系统的安装。',
    '设备安装顺序：先大型后小型、先主体后辅助、先地下后地上。',
    '管道安装严格按工艺流程图标高和走向施工，阀门安装位置应便于操作检修。',
    '电气安装满足相关规范要求，做好接地和防雷保护。',
    '施工过程中做好成品保护，防止已安装设备受到污染或损坏。',
], 1):
    add_item('%d）%s' % (i, item))

add_h3('（三）调试阶段')
for i, item in enumerate([
    '调试阶段的主要工作是将废水按工艺要求逐步接入新建系统，逐步调试至达标状态。',
    '单机试车：检查各台设备（泵、风机、搅拌机、压滤机等）运行是否正常。',
    '清水联动试车：所有设备联动运行24~48 h，检查管路、阀门、仪表。',
    '工艺调试：接种活性污泥，低负荷启动生化系统，逐步提高进水负荷至设计值。调试周期约30天。',
    '含铀废水处理系统调试期间应加强辐射监测，确保操作人员安全。',
], 1):
    add_item('%d）%s' % (i, item))

# ==================== CHAPTER 6 ====================
page_break()
add_h1('第六章  工程投资估算与运行成本')

add_h2('6.1  估算依据')
for i, ref in enumerate([
    '《内蒙古自治区建设工程工程量清单计价定额》（最新版）；',
    '《内蒙古自治区建筑工程计价表》；',
    '《内蒙古自治区安装工程计价表》；',
    '类似含铀废水处理工程技术经济指标；',
    '有关设备、仪器、仪表、膜组件的市场询价；',
    '《内蒙古自治区环境工程设计收费标准》。',
], 1):
    add_item('%d）%s' % (i, ref))

add_h2('6.2  投资估算')
add_p('项目工程分为土建及安装两部分，估算投资详见表6-1至表6-4。')

add_tbl('表6-1  土建工程投资估算',
    ['序号', '构筑物名称', '规格尺寸', '数量', '估算金额（万元）'],
    [['1', '原水调节池', '8.0 x 4.0 x 4.0 m 地下钢砼防腐', '1 座', '15.00'],
     ['2', '综合调节池', '8.0 x 6.0 x 4.0 m 地下钢砼防腐', '1 座', '20.00'],
     ['3', '设备基础及管沟', '钢砼结构', '1 套', '8.00'],
     ['4', '操作用房', '9.0 x 4.0 x 3.0 m 框架砖混', '1 座', '12.00'],
     ['', '小  计', '', '', '55.00']])

add_tbl('表6-2  设备及安装工程投资估算',
    ['序号', '设备名称', '规格型号', '数量', '金额（万元）'],
    [['1', '原水提升泵及搅拌系统', 'Q=12 m3/h，含曝气搅拌', '1 批', '1.10'],
     ['2', '絮凝沉淀池（含搅拌/斜管/堰/排泥）', '9.0x3.0x3.5 m，FRP防腐', '1 套', '16.80'],
     ['3', '加药装置（PAC/PAM/CaO）', 'V=1.5 m3，一箱二泵', '3 套', '2.25'],
     ['4', '多介质过滤器（含泵/水箱）', 'o1400x2850，2台，衬胶', '1 套', '14.24'],
     ['5', 'DTRO装置（并联）', '12 m3/h，75 bar，含高压泵/保安/清洗', '2 套', '360.00'],
     ['6', 'RO装置（并联）', '9 m3/h，一级二段，含高压泵/保安/清洗', '2 套', '100.00'],
     ['7', '除氟除硬沉淀池（2套）', '9.0x2.5x3.0 m，FRP防腐', '2 套', '29.00'],
     ['8', '脱氨塔', '6 m3/h，304不锈钢', '1 套', '80.00'],
     ['9', 'MVR蒸发装置', '5.4 m3/h，2205双相钢', '1 套', '350.00'],
     ['10', '刮板蒸发干化装置', '0.45 m3/h，2205双相钢', '1 套', '80.00'],
     ['11', '水解酸化池（含搅拌/布水/泵）', '10.0x3.0x5.0 m，FRP防腐', '1 套', '23.84'],
     ['12', '缺氧池（含搅拌）', '10.0x3.0x5.0 m，FRP防腐', '1 套', '24.80'],
     ['13', '好氧池2套（含曝气器）', '10.0x3.0x5.0 m，FRP防腐', '2 套', '48.00'],
     ['14', 'MBR膜组件', 'PTFE，240 m2/套 x 4套', '4 套', '72.00'],
     ['15', '芬顿氧化塔2套（含加药）', '9.0x3.0x5.0 m，FRP防腐', '2 套', '47.75'],
     ['16', '鼓风机（空气悬浮）', 'Q=10 m3/min，N=11 kW', '2 台', '24.00'],
     ['17', '污泥处理系统（含池/压滤/泵等）', '80 m2板框压滤机 x 2', '1 套', '50.35'],
     ['18', '铵盐沉淀及硅胶吸附系统', '含反应槽/压滤/吸附柱/泵组', '1 套', '25.00'],
     ['19', '仪表自控系统', '液位/pH/DO/PLC/电气柜', '1 批', '100.00'],
     ['20', '安装材料（电缆桥架+阀门管道）', '—', '1 批', '30.00'],
     ['21', '安装费（含吊装/差旅/调试）', '—', '1 批', '60.00'],
     ['', '小  计', '', '', '1,559.13']])

add_tbl('表6-3  工程投资汇总',
    ['序号', '费用类别', '金额（万元）', '占比（%）'],
    [['1', '土建工程', '55.00', '3.4'],
     ['2', '设备及安装工程', '1,559.13', '96.6'],
     ['', '合  计', '约 1,614', '100.0']])

add_p('工程估算总投资约1,614万元。吨水投资：1,614 / 288 ~ 5.60 万元/(m3-d)。', indent=False)
add_p('注：含铀废水处理因涉及膜系统（DTRO+RO）、蒸发系统（MVR）和特种防腐材料（FRP/ECTFE/2205），吨水投资显著高于常规工业废水处理。')

add_h2('6.3  运行成本')
add_p('项目运行成本包括：电费、药剂费、膜更换及维护费、人工费、污泥处置费、蒸汽费。')

add_h3('（1）电费')
add_p('系统总装机功率约416 kW，主要耗电设备为DTRO高压泵（~90 kW）、MVR压缩机（~180 kW）、鼓风机（11 kW）及各类水泵。运行系数取0.5~0.9，日均耗电量约8,446 kWh。包头地区工业电价约0.6元/kWh。')
add_p('日电费：8,446 x 0.6 ~ 5,068 元/天；年电费约185.0万元。', indent=False)

add_h3('（2）药剂费')
add_p('包括PAC、PAM、CaO/Ca(OH)2、Na2CO3、H2O2、FeSO4、H2SO4、NaOH、氨水、硝酸、次氯酸钠等。综合药剂费按3.5元/m3计。')
add_p('日药剂费：12 x 24 x 3.5 ~ 1,008 元/天；年药剂费约36.8万元。', indent=False)

add_h3('（3）膜更换及维护费')
add_p('DTRO膜柱使用寿命3~5年，RO膜元件3~5年，MBR膜（PTFE）5~8年。年维护费按膜系统设备费的5%计提。')
add_p('年维护费：(360 + 100 + 72) x 5% ~ 26.6万元/年。', indent=False)

add_h3('（4）人工费')
add_p('废水处理站设操作人员6人（含技术负责人1人），四班三运转。人均年工资按8万元计。')
add_p('年人工费：6 x 8 = 48.0万元/年。', indent=False)

add_h3('（5）污泥处置费')
add_p('系统日产含铀污泥（含水率约70%）约1.5吨，年产量约550吨。含铀污泥属于放射性废物，处置费用约2000元/吨（含运输）。')
add_p('年污泥处置费：550 x 2000 = 110.0万元/年。', indent=False)

add_h3('（6）蒸汽费（MVR补充）')
add_p('MVR蒸发系统启动和运行中需少量补充蒸汽（约0.5 t/h），蒸汽单价约200元/t。')
add_p('年蒸汽费：0.5 x 24 x 365 x 200 ~ 87.6万元/年。', indent=False)

add_tbl('表6-4  运行成本汇总',
    ['序号', '成本项目', '日费用（元/天）', '年费用（万元/年）', '占比（%）'],
    [['1', '电费', '5,068', '185.0', '37.5'],
     ['2', '药剂费', '1,008', '36.8', '7.4'],
     ['3', '膜更换及维护费', '729', '26.6', '5.4'],
     ['4', '人工费', '1,315', '48.0', '9.7'],
     ['5', '污泥处置费', '3,014', '110.0', '22.3'],
     ['6', '蒸汽费（MVR补充）', '2,400', '87.6', '17.7'],
     ['', '合  计', '13,534', '494.0', '100.0']])

add_p('年运行总成本约494万元/年。', indent=False)
add_p('吨水运行成本：494.0 / (12 x 24 x 365) x 10000 ~ 47.0 元/m3。', indent=False)
add_p('设备折旧（按15年，设备费约1,559万元）：1,559 / 15 ~ 103.9万元/年。', indent=False)
add_p('吨水综合成本（含折旧）：(494.0 + 103.9) / 105,120 x 10000 ~ 56.9 元/m3。', indent=False)
add_p('注：含铀废水处理运行成本较高，主要受电耗（37.5%）、污泥处置（22.3%）和MVR蒸汽补充（17.7%）三项影响。与传统工业废水生化处理（约5~10元/m3）相比，增量成本主要来自膜浓缩、蒸发结晶和放射性废物处置。')

# ==================== FOOTER ====================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '—— 方案完 ——', size=Pt(10), color=RGBColor(0x99, 0x99, 0x99))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '本方案根据《污水综合排放标准》（GB 8978-1996）、《铀矿冶辐射防护和环境保护规定》（GB 23727-2009）、', size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '《放射性废物管理规定》（GB 14500-2002）、《铀矿冶废水治理工程技术规范》（HJ 2048-2015）等编制', size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '内蒙古包头市含铀废水处理工程  |  ' + datetime.date.today().strftime('%Y年%m月'), size=Pt(9), color=RGBColor(0x99, 0x99, 0x99))

output = r'E:\claude\wastewater-app\backend\reports\baotou_uranium_flowchart.docx'
doc.save(output)
print('DOCX saved: ' + output)
